from typing import Iterator
from langchain_core.documents import Document
from langchain_core.document_loaders import BaseLoader
from docx import Document as DocxDocument

class CustomDocLoader(BaseLoader):
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path

    def lazy_load(self) -> Iterator[Document]:
        doc = DocxDocument(self.file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        yield Document(
            page_content="\n".join(full_text),
            metadata={"source": self.file_path}
        )

if __name__ == "__main__":
    import os
    absfile = os.path.abspath(__file__)
    dirfile = os.path.dirname(absfile)
    dirfile = os.path.dirname(dirfile)
    file = os.path.join(dirfile, 'data/milvus向量数据库笔记2.docx')
    loader = CustomDocLoader(file)
    for doc in loader.lazy_load():
        print(doc.page_content[:100])  # 打印前100个字符，验证加载是否成功
        print(doc.metadata.get("source"))  # 打印元数据，验证文件路径是否正确